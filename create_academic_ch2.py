from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_math(p, text):
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    return p

title = doc.add_heading('', level=0)
run = title.add_run('Chương 2: Cơ sở Lý thuyết (Theoretical Foundations)')
run.font.size = Pt(16)
run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('2.1. Hình thức hóa mô hình Retrieval-Augmented Generation', level=1)

doc.add_heading('2.1.1. Kiến trúc không gian vector mật độ cao (Dense Vector Space)', level=2)
p = doc.add_paragraph('Về mặt toán học, hệ thống RAG tiêu chuẩn (Naive RAG) hoạt động thông qua một hàm nhúng (embedding function) ')
add_math(p, 'Φ: X → R^d')
p.add_run(', trong đó ')
add_math(p, 'X')
p.add_run(' là không gian rời rạc của các văn bản và ')
add_math(p, 'R^d')
p.add_run(' là không gian liên tục d-chiều. Cho ngữ liệu ')
add_math(p, 'C = {c_1, c_2, ..., c_N}')
p.add_run(' và truy vấn ')
add_math(p, 'q')
p.add_run(', hệ thống tối đa hóa hàm tương đồng Cosine:')

p = doc.add_paragraph()
add_math(p, 'S_C(q, c_i) = (Φ(q) · Φ(c_i)) / (||Φ(q)||₂ × ||Φ(c_i)||₂)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Hàm này chiếu các văn bản lên bề mặt của một siêu cầu (hypersphere) ')
add_math(p, 'S^{d-1}')
p.add_run(', nơi khoảng cách được đo bằng góc ')
add_math(p, 'θ(u, v) = arccos(S_C(u, v))')
p.add_run('.')

doc.add_heading('2.1.2. Lời nguyền chiều dữ liệu và sự sụp đổ tính bắc cầu (Transitivity Collapse)', level=2)
p = doc.add_paragraph('Bất chấp hiệu quả trong việc tìm kiếm ngữ nghĩa đơn lẻ, Dense Retrieval đối mặt với giới hạn toán học cốt lõi khi xử lý suy luận đa bước (multi-hop reasoning).')

p = doc.add_paragraph('Suy luận đa bước yêu cầu đi qua các chuỗi logic (ví dụ: A → B → C). Trên không gian vector, khoảng cách góc tuân thủ chặt chẽ bất đẳng thức tam giác:')

p = doc.add_paragraph()
add_math(p, 'θ(A, C) ≤ θ(A, B) + θ(B, C)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Nếu A và C không có điểm chung từ vựng trực tiếp nhưng liên kết qua B, các vector của chúng thường trực giao (orthogonal, ')
add_math(p, 'θ(A, C) ≈ π/2')
p.add_run('). Việc áp dụng cộng tuyến tính trên bề mặt đa tạp của ngôn ngữ tự nhiên thường thất bại, khiến việc truy xuất qua các "nút thắt" trung gian rơi vào trạng thái nhiễu ngẫu nhiên. Lỗi này không phải do mô hình yếu, mà là hạn chế cố hữu của hình học không gian phẳng.')

doc.add_heading('2.2. Đồ thị Tri thức và Định lý Cấu trúc (Knowledge Graphs Topology)', level=1)

doc.add_heading('2.2.1. Hình thức hóa Đồ thị Thuộc tính (Property Graphs)', level=2)
p = doc.add_paragraph('Khác với các RDF Triplestore (Bộ ba Subject-Predicate-Object) truyền thống, GraphRAG sử dụng Đồ thị Thuộc tính có độ biểu đạt cao hơn, được định nghĩa là một hệ:')

p = doc.add_paragraph()
add_math(p, 'G = (V, E, ρ, λ, σ)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Trong đó:')
p = doc.add_paragraph('• V là tập hợp hữu hạn các đỉnh (thực thể).')
p = doc.add_paragraph('• E là tập hợp hữu hạn các cạnh có hướng (mối quan hệ).')
p = doc.add_paragraph('• ρ: E → V × V là hàm tỷ lệ ánh xạ cạnh tới đỉnh nguồn và đích.')
p = doc.add_paragraph('• λ: (V ∪ E) → Σ gán nhãn phân loại từ tập Σ cho đỉnh và cạnh.')
p = doc.add_paragraph('• σ: (V ∪ E) × K → W là hàm ánh xạ từng phần gán thuộc tính, biến đỉnh/cạnh và khóa k ∈ K thành giá trị w ∈ W (đặc biệt là lưu trữ mô tả văn bản tự nhiên).')

doc.add_heading('2.2.2. Toán học của Personalized PageRank (PPR) trong truy xuất', level=2)
p = doc.add_paragraph('Để khắc phục sự cố góc trong vector, không gian đồ thị sử dụng cấu trúc tô-pô. Các biến thể của GraphRAG tính toán độ quan trọng của nút dựa trên Personalized PageRank (PPR).')

p = doc.add_paragraph('Cho đồ thị có N đỉnh, ma trận kề A, và ma trận đường chéo bậc D. Ma trận chuyển đổi xác suất là ')
add_math(p, 'P = A D^{-1}')
p.add_run('. Đối với truy vấn tại nút nguồn s (với vector cơ sở e_s), phân phối dừng π_s của PPR thỏa mãn phương trình:')

p = doc.add_paragraph()
add_math(p, 'π_s = (1 - α) P π_s + α e_s')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Trong đó α ∈ [0.15, 0.3] là xác suất khởi động lại (restart probability). Lời giải giải tích đóng của chuỗi Markov này là:')

p = doc.add_paragraph()
add_math(p, 'π_s = α [I - (1 - α) P]^{-1} e_s = α Σ_{k=0}^∞ (1 - α)^k P^k e_s')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Việc khai triển chuỗi Neumann ở vế phải chính là chứng minh toán học của multi-hop reasoning. Mỗi số hạng ')
add_math(p, 'P^k e_s')
p.add_run(' biểu diễn xác suất chạm đến đích sau chính xác k bước nhảy. Cấu trúc này cho phép suy luận bắc cầu mà không làm mất đi định hướng ngữ nghĩa như trong tính toán Cosine.')

doc.add_heading('2.3. Tối ưu hóa Tính Mô-đun và Thuật toán Leiden', level=1)

doc.add_heading('2.3.1. Bài toán phát hiện cộng đồng và Modularity (Q)', level=2)
p = doc.add_paragraph('Để thực hiện bài toán tổng hợp toàn cục (Global Sensemaking), GraphRAG áp dụng lý thuyết mạng để nhóm các đỉnh thành các cộng đồng C = {C_1, ..., C_k}. Hàm mục tiêu cần tối đa hóa là Modularity (Tính mô-đun):')

p = doc.add_paragraph()
add_math(p, 'Q = (1/2m) Σ_{i,j} [A_{ij} - (k_i k_j)/(2m)] δ(c_i, c_j)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Trong đó, m là tổng trọng số các cạnh, k_i là bậc của đỉnh i. Đại lượng ')
add_math(p, '(k_i k_j)/(2m)')
p.add_run(' là số cạnh kỳ vọng giữa đỉnh i và j trong một đồ thị ngẫu nhiên. Bằng cách lấy phân phối thực tế trừ đi kỳ vọng ngẫu nhiên, hệ thống tìm ra những cụm có mật độ kết nối vượt trội.')

doc.add_heading('2.3.2. Sự sụp đổ của Louvain và giải pháp Leiden', level=2)
p = doc.add_paragraph('Thuật toán Louvain, tiền thân của thuật toán phát hiện cộng đồng, hoạt động theo nguyên tắc gộp tham lam. Tuy nhiên, định lý toán học chỉ ra Louvain mắc lỗi "Disconnected Community Problem" - tạo ra các cộng đồng mà bên trong chúng bị ngắt kết nối (lên tới 25% số cộng đồng).')

p = doc.add_paragraph('GraphRAG (Edge et al., 2024) khắc phục triệt để điều này bằng thuật toán Leiden (Traag et al., 2019), bao gồm ba pha nghiêm ngặt:')
p = doc.add_paragraph('1. Local Moving: Di chuyển nút để tối đa hóa sự thay đổi ΔQ.')
p = doc.add_paragraph('2. Refinement (Điểm khác biệt lõi): Các nút trong cộng đồng được tách ra thành các nút đơn lẻ, và chỉ được gộp lại nếu thỏa mãn ràng buộc liên thông chặt chẽ bên trong ranh giới cộng đồng gốc.')
p = doc.add_paragraph('3. Aggregation: Xây dựng đồ thị quy mô lớn hơn từ các cụm đã tinh chỉnh.')

p = doc.add_paragraph('Nhờ bổ sung pha Tinh chỉnh, Leiden đảm bảo 100% các cộng đồng tạo ra đều mang tính liên thông yếu (weakly connected), cung cấp một bộ tóm tắt có ngữ cảnh mạch lạc tuyệt đối cho LLM trong giai đoạn Map-Reduce sau này.')

doc.save('Chapter2_Academic.docx')
print('Academic Chapter 2 written successfully.')
