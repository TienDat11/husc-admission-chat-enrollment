import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_doc(md_file_path, output_docx_path):
    # Initialize Word Document
    doc = Document()
    
    # Configure default styles to make it look academic
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    
    # Set Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Set Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)

    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Pre-process math equations to text logic
    # Math replacement patterns (Basic translation to human-readable text)
    math_replacements = {
        r'\$\\mathcal\{C\} = \\\{d_1, d_2, \.\.\., d_n\\\}\$': 'tập ngữ liệu văn bản đầu vào C, bao gồm các tài liệu d_1, d_2, ..., d_n',
        r'\$d_i\$': 'tài liệu thứ i',
        r'\$\\mathcal\{C\}\$': 'tập ngữ liệu C',
        r'\$G = \(V, E\)\$': 'đồ thị tri thức G gồm các đỉnh V và các cạnh E',
        r'\$V = \\\{v_1, v_2, \.\.\., v_N\\\}\$': 'tập đỉnh V gồm các thực thể v_1, v_2, ..., v_N',
        r'\$v_i\$': 'đỉnh thứ i',
        r'\$P\(v_i\)\$': 'tập thuộc tính P của đỉnh thứ i',
        r'\$E \\subseteq V \\times V\$': 'tập cạnh E biểu diễn mối liên hệ giữa các đỉnh trong V',
        r'\$e_\{ij\} \\in E\$': 'cạnh e_ij thuộc tập E',
        r'\$v_i\$ và \$v_j\$': 'đỉnh thứ i và đỉnh thứ j',
        r'\$w_\{ij\}\$': 'trọng số w_ij',
        r'\$\\mathcal\{P\} = \\\{C_1, C_2, \.\.\., C_k\\\}\$': 'tập hợp các cộng đồng P bao gồm C_1, C_2, ..., C_k',
        r'\$C_i\$': 'cộng đồng thứ i',
        r'\$E_\{in\}\$': 'mật độ liên kết nội bộ E_in',
        r'\$V\$': 'tập đỉnh V',
        r'\$G\$': 'đồ thị G',
        r'\$E\$': 'tập cạnh E',
        r'\$\\mathcal\{P\}\$': 'tập hợp các cộng đồng P',
        r'\$q\$': 'câu hỏi đầu vào q'
    }

    for pattern, replacement in math_replacements.items():
        content = re.sub(pattern, replacement, content)

    # Process line by line
    lines = content.split('\n')
    
    title = lines[0].replace('# ', '').strip()
    
    # Add Title
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    doc.add_paragraph() # Empty line

    i = 1
    while i < len(lines):
        line = lines[i].strip()
        
        if not line or line == '---':
            i += 1
            continue
            
        if line.startswith('## '):
            heading = line.replace('## ', '')
            doc.add_heading(heading, level=1)
        elif line.startswith('### '):
            heading = line.replace('### ', '')
            doc.add_heading(heading, level=2)
        elif line.startswith('**') and line.endswith('**') and len(line) < 100:
            doc.add_heading(line.replace('**', ''), level=2)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Bold parsing within list
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if j % 2 == 1: # bold
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
            text = line[3:]
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if j % 2 == 1:
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)
        else:
            # Normal paragraph
            # Basic bold parsing
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for j, part in enumerate(parts):
                if j % 2 == 1:
                    p.add_run(part).bold = True
                else:
                    # Basic italic parsing
                    sub_parts = re.split(r'\*(.*?)\*', part)
                    for k, sub_part in enumerate(sub_parts):
                        if k % 2 == 1:
                            p.add_run(sub_part).italic = True
                        else:
                            p.add_run(sub_part)

        # Image insertion logic based on context
        if "Naive RAG và Dense Retrieval" in line:
            # Insert diagram 1 after section 2.1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture('diagram1.png', width=Inches(6.0))
                p_cap = doc.add_paragraph('Hình 1: So sánh cấu trúc xử lý giữa Naive RAG và GraphRAG', style='Caption')
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Lỗi chèn ảnh: {e}")

        if "Global Search (Truy vấn toàn cục" in line:
            # Insert diagram 2 after Global Search context
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture('diagram2.png', width=Inches(5.0))
                p_cap = doc.add_paragraph('Hình 2: Kiến trúc truy vấn toàn cục sử dụng cơ chế Map-Reduce', style='Caption')
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Lỗi chèn ảnh: {e}")

        i += 1

    doc.save(output_docx_path)
    print(f"Document saved to {output_docx_path}")

if __name__ == '__main__':
    create_word_doc('GraphRAG_Research_Document.md', 'GraphRAG_Paper.docx')
