from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_math(p, text):
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    return p

title = doc.add_heading('', level=0)
run = title.add_run('Chương 2: Cơ sở Lý thuyết (Theoretical Foundations)')
run.font.size = Pt(16)
run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 2.1
doc.add_heading('2.1. Hình thức hóa mô hình Retrieval-Augmented Generation', level=1)

doc.add_heading('2.1.1. Kiến trúc không gian vector mật độ cao (Dense Vector Space)', level=2)
p = doc.add_paragraph('Về mặt toán học, hệ thống RAG tiêu chuẩn hoạt động thông qua một hàm nhúng (embedding function) ')
add_math(p, 'Φ: X → R^d')
p.add_run(', trong đó ')
add_math(p, 'X')
p.add_run(' là không gian rời rạc của các đoạn văn bản (text chunks) và ')
add_math(p, 'R^d')
p.add_run(' là không gian vector liên tục d-chiều. Cho trước một ngữ liệu ')
add_math(p, 'C = {c_1, c_2, ..., c_N}')
p.add_run(' và một truy vấn ')
add_math(p, 'q')
p.add_run(', mục tiêu của hệ thống là tối đa hóa hàm mục tiêu tương đồng:')

p = doc.add_paragraph()
add_math(p, 'S(q, c_i) = cos(Φ(q), Φ(c_i)) = (Φ(q) · Φ(c_i)) / (||Φ(q)||₂ × ||Φ(c_i)||₂)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Các mô hình nhúng hiện đại (như text-embedding-3-small hay BGE-M3) được huấn luyện thông qua hàm mất mát đối nghịch (Contrastive Loss) như InfoNCE để đẩy các vector ngữ nghĩa tương đồng lại gần nhau trên bề mặt của một siêu cầu (hypersphere) có chuẩn bằng 1.')

doc.add_heading('2.1.2. Giới hạn tô-pô của tìm kiếm vector (Topological Limitations)', level=2)
p = doc.add_paragraph('Sự thất bại của Dense Retrieval đối với các truy vấn tổng hợp (sensemaking queries) bắt nguồn từ một nghịch lý cấu trúc cơ bản. Việc chia nhỏ văn bản dài thành các chunk ')
add_math(p, 'c_i')
p.add_run(' độc lập phá vỡ đồ thị phụ thuộc (dependency graph) tự nhiên của văn bản. Khoảng cách Euclid hay Cosine Similarity chỉ đo lường được sự tương đồng về mặt phân phối từ vựng (distributional semantics), nhưng hoàn toàn vô lực trong việc nắm bắt khoảng cách quan hệ (relational distance).')

p = doc.add_paragraph('Xét một bài toán multi-hop reasoning, nơi thực thể ')
add_math(p, 'v_A')
p.add_run(' liên hệ với ')
add_math(p, 'v_B')
p.add_run(' trong đoạn văn ')
add_math(p, 'c_1')
p.add_run(', và ')
add_math(p, 'v_B')
p.add_run(' liên hệ với ')
add_math(p, 'v_C')
p.add_run(' trong đoạn văn ')
add_math(p, 'c_2')
p.add_run('. Bất đẳng thức tam giác trong không gian metrid cho vector không đảm bảo rằng vector của truy vấn chứa ')
add_math(p, 'v_A')
p.add_run(' sẽ đủ gần vector của ')
add_math(p, 'c_2')
p.add_run(' (chứa ')
add_math(p, 'v_C')
p.add_run('). Hậu quả là chuỗi liên kết suy luận bị đứt gãy.')

# 2.2
doc.add_heading('2.2. Đồ thị Tri thức (Knowledge Graphs)', level=1)

doc.add_heading('2.2.1. Định nghĩa toán học của Đồ thị Thuộc tính (Property Graphs)', level=2)
p = doc.add_paragraph('Thay vì biểu diễn dạng bộ ba (triples) truyền thống trong RDF (Resource Description Framework), GraphRAG áp dụng cấu trúc Đồ thị Thuộc tính có hướng (Directed Property Graph), được định nghĩa là một tuple bộ sáu:')

p = doc.add_paragraph()
add_math(p, 'G = (V, E, L_V, L_E, P_V, P_E)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Trong đó:')
p = doc.add_paragraph('• ')
add_math(p, 'V')
p.add_run(' là tập hợp hữu hạn các đỉnh (thực thể).')
p = doc.add_paragraph('• ')
add_math(p, 'E ⊆ V × V')
p.add_run(' là tập hợp các cạnh (mối quan hệ).')
p = doc.add_paragraph('• ')
add_math(p, 'L_V')
p.add_run(' và ')
add_math(p, 'L_E')
p.add_run(' là các hàm gán nhãn phân loại cho đỉnh và cạnh.')
p = doc.add_paragraph('• ')
add_math(p, 'P_V')
p.add_run(' và ')
add_math(p, 'P_E')
p.add_run(' là các hàm ánh xạ gán tập thuộc tính (properties) dưới dạng key-value. Đặc biệt trong GraphRAG, thuộc tính quan trọng nhất là các chuỗi văn bản tự nhiên mô tả chi tiết thực thể và quan hệ.')

doc.add_heading('2.2.2. Khoảng cách đồ thị và Suy luận đa bước (Graph Distance and Multi-hop Reasoning)', level=2)
p = doc.add_paragraph('Trên không gian đồ thị, hàm khoảng cách ')
add_math(p, 'd_G(v_i, v_j)')
p.add_run(' không được tính bằng khoảng cách cosine, mà được tính thông qua độ dài đường đi ngắn nhất (shortest path) hoặc các thuật toán đi dạo ngẫu nhiên (Random Walk) như Personalized PageRank (PPR). Sự chuyển đổi hệ quy chiếu từ "khoảng cách hình học" sang "khoảng cách tô-pô" cho phép mô hình truy vết các quan hệ gián tiếp một cách tường minh và mang tính giải thích cao (explainable).')

# 2.3
doc.add_heading('2.3. Tối ưu hóa Tính Mô-đun và Thuật toán Leiden', level=1)

doc.add_heading('2.3.1. Bài toán Phân cụm Đồ thị (Graph Partitioning)', level=2)
p = doc.add_paragraph('Để tóm tắt đồ thị quy mô lớn (Graph Summarization), đồ thị cần được chia nhỏ thành các cụm thông tin có ý nghĩa. Gọi tập các cộng đồng phân hoạch của đồ thị là ')
add_math(p, 'C = {C_1, C_2, ..., C_k}')
p.add_run(', với điều kiện ')
add_math(p, 'C_i ∩ C_j = ∅')
p.add_run(' và ')
add_math(p, '∪ C_i = V')
p.add_run('. Mục tiêu là tìm một phân hoạch tối đa hóa mật độ cạnh bên trong cộng đồng và tối thiểu hóa mật độ cạnh giữa các cộng đồng.')

doc.add_heading('2.3.2. Hàm mục tiêu Modularity (Q)', level=2)
p = doc.add_paragraph('Độ đo tiêu chuẩn cho bài toán này là Modularity, được giới thiệu bởi Newman (2006). Đối với đồ thị có trọng số, nó được định nghĩa là:')

p = doc.add_paragraph()
add_math(p, 'Q = (1/2m) × Σ_{i,j} [A_{ij} - (k_i × k_j)/(2m)] × δ(c_i, c_j)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Giải thích các thành phần:')
p = doc.add_paragraph('• ')
add_math(p, 'A_{ij}')
p.add_run(' là trọng số cạnh giữa đỉnh i và j trong ma trận kề.')
p = doc.add_paragraph('• ')
add_math(p, 'k_i = Σ_j A_{ij}')
p.add_run(' là bậc (tổng trọng số) của đỉnh i.')
p = doc.add_paragraph('• ')
add_math(p, 'm = (1/2) Σ_{i,j} A_{ij}')
p.add_run(' là tổng trọng số của toàn bộ đồ thị.')
p = doc.add_paragraph('• ')
add_math(p, '(k_i × k_j)/(2m)')
p.add_run(' là giá trị kỳ vọng (expected weight) của cạnh nối i và j trong một đồ thị ngẫu nhiên (Null model) có cùng phân phối bậc.')
p = doc.add_paragraph('• ')
add_math(p, 'δ(c_i, c_j)')
p.add_run(' là hàm Kronecker delta, bằng 1 nếu đỉnh i và j thuộc cùng một cộng đồng (c_i = c_j), ngược lại bằng 0.')

doc.add_heading('2.3.3. Hạn chế của Thuật toán Louvain', level=2)
p = doc.add_paragraph('Thuật toán Louvain áp dụng chiến lược tham lam (greedy) để tối ưu hóa Q. Tuy nhiên, Traag et al. (2019) đã chứng minh bằng toán học rằng Louvain có một lỗ hổng nghiêm trọng: nó có thể tạo ra các cộng đồng bị đứt gãy nội bộ (internally disconnected). Một đỉnh có thể đóng vai trò "cầu nối" trong một cộng đồng, khi đỉnh đó bị di chuyển sang cộng đồng khác, phần còn lại của cộng đồng ban đầu bị vỡ thành hai mảnh không liên thông. Thực nghiệm cho thấy có đến 25% cộng đồng do Louvain sinh ra gặp phải các vấn đề kết nối kém.')

doc.add_heading('2.3.4. Sự bảo đảm kết nối của Thuật toán Leiden', level=2)
p = doc.add_paragraph('GraphRAG thay thế Louvain bằng thuật toán Leiden. Leiden khắc phục triệt để lỗi phân mảnh bằng cách thêm một bước Refinement (Tinh chỉnh) cực kỳ chặt chẽ. Cụ thể, sau pha di chuyển cục bộ ban đầu (Local Moving), Leiden không gộp ngay các cộng đồng. Thay vào đó, nó chia mỗi cộng đồng thành các nút đơn lẻ, sau đó cho phép các nút này gộp lại với nhau (có chọn lọc) chỉ khi chúng bị ràng buộc nghiêm ngặt trong ranh giới của cộng đồng gốc.')

p = doc.add_paragraph('Về mặt toán học, thuật toán Leiden đưa ra hai định lý quan trọng:')
p = doc.add_paragraph('1. Định lý Kết nối (Connectivity Theorem): Mọi cộng đồng được sinh ra bởi Leiden đảm bảo tính liên thông yếu (weakly connected).')
p = doc.add_paragraph('2. Định lý Cận dưới (Lower Bound Theorem): Modularity đầu ra của Leiden luôn lớn hơn hoặc bằng kết quả của thuật toán Louvain.')

p = doc.add_paragraph('Thuật toán Leiden kết xuất ra cấu trúc đồ thị phân cấp (Hierarchical Graph). Lớp gốc rễ bao gồm một vài cụm rất lớn (C0 - Macro level) và phân tách dần thành các cụm nhỏ hơn (C1, C2, C3 - Micro level). Bằng cách để LLM sinh tóm tắt cho từng mức độ cộng đồng, GraphRAG chính thức giải quyết được bài toán tổng hợp thông tin đa tầng một cách hoàn chỉnh.')

doc.save('Chapter2_Advanced.docx')
print('Advanced Chapter 2 completed')
